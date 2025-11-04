#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
文件分析工具 - 露尼西亚的文件处理能力
支持PDF、CSV、Excel等文件类型的智能分析
使用LangChain作为统一接口，PyMuPDF处理PDF，pandas处理表格
"""

import os
import json
import fitz  # PyMuPDF
import pandas as pd
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass
import io
import base64
from pathlib import Path
import sys

# 安全打印函数（避免Windows终端emoji编码问题）
def safe_print(msg):
    """安全打印，避免emoji编码错误"""
    try:
        print(msg)
    except UnicodeEncodeError:
        # 移除所有非ASCII和常见中文之外的字符（包括emoji）
        import re
        # 只保留ASCII、中文、常见标点
        msg_safe = re.sub(r'[^\x00-\x7F\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]', '', msg)
        try:
            print(msg_safe)
        except UnicodeEncodeError:
            # 最终回退：只保留ASCII
            msg_ascii = re.sub(r'[^\x00-\x7F]', '', msg)
            print(msg_ascii)

# python-docx for Word documents
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    safe_print("⚠️ python-docx 未安装，Word文档分析功能不可用")

# Code analyzer for programming languages
from code_analyzer import PythonCodeAnalyzer, GeneralCodeAnalyzer, CodeAnalysisResult

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.document_loaders.base import BaseLoader

@dataclass
class FileAnalysisResult:
    """文件分析结果"""
    file_type: str
    file_name: str
    content: str
    metadata: Dict[str, Any]
    summary: str
    analysis: str
    success: bool
    error: Optional[str] = None

class PDFAnalyzer:
    """PDF文件分析器"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
    
    def extract_text(self, file_path: str) -> FileAnalysisResult:
        """提取PDF文本内容"""
        try:
            safe_print(f"📄 开始分析PDF文件: {file_path}")
            
            # 使用PyMuPDF打开PDF
            doc = fitz.open(file_path)
            text_content = ""
            metadata = {
                "page_count": doc.page_count,
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "creator": doc.metadata.get("creator", ""),
                "producer": doc.metadata.get("producer", ""),
                "creation_date": doc.metadata.get("creationDate", ""),
                "modification_date": doc.metadata.get("modDate", ""),
            }
            
            # 提取所有页面的文本
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                text_content += f"\n--- 第{page_num + 1}页 ---\n{page_text}"
            
            doc.close()
            
            # 使用LangChain分割文本
            documents = self.text_splitter.split_text(text_content)
            
            # 生成摘要和分析
            summary = self._generate_summary(text_content, metadata)
            analysis = self._analyze_content(text_content, metadata)
            
            return FileAnalysisResult(
                file_type="PDF",
                file_name=os.path.basename(file_path),
                content=text_content,
                metadata=metadata,
                summary=summary,
                analysis=analysis,
                success=True
            )
            
        except Exception as e:
            safe_print(f"❌ PDF分析失败: {e}")
            return FileAnalysisResult(
                file_type="PDF",
                file_name=os.path.basename(file_path),
                content="",
                metadata={},
                summary="",
                analysis="",
                success=False,
                error=str(e)
            )
    
    def _generate_summary(self, content: str, metadata: Dict) -> str:
        """生成PDF摘要"""
        lines = content.split('\n')
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        
        summary_parts = []
        
        # 基本信息
        if metadata.get("title"):
            summary_parts.append(f"📖 标题: {metadata['title']}")
        if metadata.get("author"):
            summary_parts.append(f"👤 作者: {metadata['author']}")
        if metadata.get("page_count"):
            summary_parts.append(f"📄 页数: {metadata['page_count']}")
        
        # 内容概览
        summary_parts.append(f"📝 内容长度: {len(content)} 字符")
        summary_parts.append(f"📊 有效行数: {len(non_empty_lines)} 行")
        
        # 关键词提取（简单版本）
        words = content.lower().split()
        word_freq = {}
        for word in words:
            if len(word) > 3:  # 只统计长度大于3的词
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # 获取高频词
        top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:5]
        if top_words:
            summary_parts.append(f"🔑 关键词: {', '.join([word for word, freq in top_words])}")
        
        return "\n".join(summary_parts)
    
    def _analyze_content(self, content: str, metadata: Dict) -> str:
        """分析PDF内容"""
        analysis_parts = []
        
        # 文档结构分析
        lines = content.split('\n')
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        
        # 检测可能的章节
        chapter_indicators = ['第', '章', 'Chapter', 'Section', '部分']
        chapters = [line for line in non_empty_lines if any(indicator in line for indicator in chapter_indicators)]
        
        if chapters:
            analysis_parts.append(f"📚 检测到 {len(chapters)} 个可能的章节标题")
        
        # 检测表格（简单检测）
        table_indicators = ['|', '\t', '  ']
        table_lines = [line for line in non_empty_lines if any(indicator in line for indicator in table_indicators)]
        if len(table_lines) > 5:
            analysis_parts.append("📊 检测到可能的表格数据")
        
        # 检测列表
        list_indicators = ['•', '-', '*', '1.', '2.', '3.']
        list_lines = [line for line in non_empty_lines if any(line.strip().startswith(indicator) for indicator in list_indicators)]
        if list_lines:
            analysis_parts.append(f"📋 检测到 {len(list_lines)} 个列表项")
        
        # 检测数字和统计信息
        import re
        numbers = re.findall(r'\d+', content)
        if len(numbers) > 10:
            analysis_parts.append("📈 包含大量数字数据，可能是统计报告")
        
        return "\n".join(analysis_parts) if analysis_parts else "📄 标准文档内容"

class TableAnalyzer:
    """表格文件分析器（CSV、Excel）"""
    
    def __init__(self):
        self.supported_formats = ['.csv', '.xlsx', '.xls']
    
    def analyze_table(self, file_path: str) -> FileAnalysisResult:
        """分析表格文件"""
        try:
            safe_print(f"📊 开始分析表格文件: {file_path}")
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # 根据文件类型读取数据
            if file_ext == '.csv':
                df = pd.read_csv(file_path, encoding='utf-8')
            elif file_ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
            
            # 生成元数据
            metadata = {
                "file_type": file_ext,
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "data_types": df.dtypes.to_dict(),
                "memory_usage": df.memory_usage(deep=True).sum(),
            }
            
            # 生成内容摘要
            content = self._generate_table_content(df)
            summary = self._generate_table_summary(df, metadata)
            analysis = self._analyze_table_data(df, metadata)
            
            return FileAnalysisResult(
                file_type="TABLE",
                file_name=os.path.basename(file_path),
                content=content,
                metadata=metadata,
                summary=summary,
                analysis=analysis,
                success=True
            )
            
        except Exception as e:
            safe_print(f"❌ 表格分析失败: {e}")
            return FileAnalysisResult(
                file_type="TABLE",
                file_name=os.path.basename(file_path),
                content="",
                metadata={},
                summary="",
                analysis="",
                success=False,
                error=str(e)
            )
    
    def _generate_table_content(self, df: pd.DataFrame) -> str:
        """生成表格内容文本"""
        content_parts = []
        
        # 添加列名
        content_parts.append("列名: " + ", ".join(df.columns.tolist()))
        content_parts.append("")
        
        # 添加前几行数据作为示例
        content_parts.append("数据预览:")
        content_parts.append(df.head(10).to_string())
        
        # 添加统计信息
        if not df.empty:
            content_parts.append("\n统计信息:")
            content_parts.append(df.describe().to_string())
        
        return "\n".join(content_parts)
    
    def _generate_table_summary(self, df: pd.DataFrame, metadata: Dict) -> str:
        """生成表格摘要"""
        summary_parts = []
        
        summary_parts.append(f"📊 数据维度: {metadata['rows']} 行 × {metadata['columns']} 列")
        summary_parts.append(f"📝 列名: {', '.join(metadata['column_names'])}")
        
        # 数据类型统计
        type_counts = {}
        for dtype in df.dtypes:
            type_name = str(dtype)
            type_counts[type_name] = type_counts.get(type_name, 0) + 1
        
        type_info = ", ".join([f"{dtype}({count})" for dtype, count in type_counts.items()])
        summary_parts.append(f"🔢 数据类型: {type_info}")
        
        # 缺失值统计
        missing_values = df.isnull().sum()
        if missing_values.sum() > 0:
            missing_info = ", ".join([f"{col}({count})" for col, count in missing_values.items() if count > 0])
            summary_parts.append(f"⚠️ 缺失值: {missing_info}")
        else:
            summary_parts.append("✅ 无缺失值")
        
        return "\n".join(summary_parts)
    
    def _analyze_table_data(self, df: pd.DataFrame, metadata: Dict) -> str:
        """分析表格数据"""
        analysis_parts = []
        
        # 数值列分析
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            analysis_parts.append(f"📈 数值列: {len(numeric_cols)} 个")
            
            # 检测异常值
            for col in numeric_cols:
                Q1 = df[col].quantile(0.25)
                Q3 = df[col].quantile(0.75)
                IQR = Q3 - Q1
                outliers = df[(df[col] < Q1 - 1.5 * IQR) | (df[col] > Q3 + 1.5 * IQR)]
                if len(outliers) > 0:
                    analysis_parts.append(f"⚠️ {col}列有{len(outliers)}个异常值")
        
        # 文本列分析
        text_cols = df.select_dtypes(include=['object']).columns
        if len(text_cols) > 0:
            analysis_parts.append(f"📝 文本列: {len(text_cols)} 个")
            
            # 检测重复值
            for col in text_cols:
                duplicates = df[col].duplicated().sum()
                if duplicates > 0:
                    analysis_parts.append(f"🔄 {col}列有{duplicates}个重复值")
        
        # 检测可能的ID列
        id_candidates = []
        for col in df.columns:
            if df[col].nunique() == len(df) and df[col].dtype == 'object':
                id_candidates.append(col)
        
        if id_candidates:
            analysis_parts.append(f"🆔 可能的ID列: {', '.join(id_candidates)}")
        
        # 检测时间列
        time_candidates = []
        for col in df.columns:
            if 'date' in col.lower() or 'time' in col.lower():
                time_candidates.append(col)
        
        if time_candidates:
            analysis_parts.append(f"⏰ 可能的时间列: {', '.join(time_candidates)}")
        
        return "\n".join(analysis_parts) if analysis_parts else "📊 标准表格数据"

class DocxAnalyzer:
    """Word文档分析器 (使用python-docx)"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
    
    def extract_text(self, file_path: str) -> FileAnalysisResult:
        """提取Word文档文本内容"""
        if not DOCX_AVAILABLE:
            return FileAnalysisResult(
                file_type="DOCX",
                file_name=os.path.basename(file_path),
                content="",
                metadata={},
                summary="",
                analysis="",
                success=False,
                error="python-docx未安装，请运行: pip install python-docx"
            )
        
        try:
            safe_print(f"📝 开始分析Word文档: {file_path}")
            
            # 使用python-docx打开Word文档
            doc = DocxDocument(file_path)
            
            # 提取元数据
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
            
            # 提取段落文本
            text_content = ""
            paragraph_count = 0
            
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:  # 只添加非空段落
                    text_content += para_text + "\n\n"
                    paragraph_count += 1
            
            # 提取表格内容
            table_content = ""
            for table_idx, table in enumerate(doc.tables):
                table_content += f"\n--- 表格 {table_idx + 1} ---\n"
                for row_idx, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_content += " | ".join(row_data) + "\n"
                table_content += "\n"
            
            # 合并所有内容
            full_content = text_content
            if table_content:
                full_content += "\n📊 文档中的表格:\n" + table_content
            
            metadata["has_tables"] = len(doc.tables) > 0
            metadata["has_images"] = self._count_images(doc)
            
            # 生成摘要和分析
            summary = self._generate_summary(text_content, metadata)
            analysis = self._analyze_content(text_content, metadata, doc)
            
            return FileAnalysisResult(
                file_type="DOCX",
                file_name=os.path.basename(file_path),
                content=full_content,
                metadata=metadata,
                summary=summary,
                analysis=analysis,
                success=True
            )
            
        except Exception as e:
            safe_print(f"❌ Word文档分析失败: {e}")
            import traceback
            traceback.print_exc()
            return FileAnalysisResult(
                file_type="DOCX",
                file_name=os.path.basename(file_path),
                content="",
                metadata={},
                summary="",
                analysis="",
                success=False,
                error=str(e)
            )
    
    def _count_images(self, doc) -> int:
        """统计文档中的图片数量"""
        try:
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            return image_count
        except:
            return 0
    
    def _generate_summary(self, content: str, metadata: Dict) -> str:
        """生成Word文档摘要"""
        summary_parts = []
        
        # 基本信息
        if metadata.get("title"):
            summary_parts.append(f"📖 标题: {metadata['title']}")
        if metadata.get("author"):
            summary_parts.append(f"👤 作者: {metadata['author']}")
        if metadata.get("subject"):
            summary_parts.append(f"📌 主题: {metadata['subject']}")
        
        # 文档结构
        summary_parts.append(f"📄 段落数: {metadata.get('paragraph_count', 0)}")
        if metadata.get("table_count", 0) > 0:
            summary_parts.append(f"📊 表格数: {metadata['table_count']}")
        if metadata.get("has_images"):
            summary_parts.append(f"🖼️ 包含图片")
        
        # 内容统计
        summary_parts.append(f"📝 内容长度: {len(content)} 字符")
        
        # 关键词
        if metadata.get("keywords"):
            summary_parts.append(f"🔑 关键词: {metadata['keywords']}")
        
        # 时间信息
        if metadata.get("created"):
            summary_parts.append(f"📅 创建时间: {metadata['created']}")
        if metadata.get("modified"):
            summary_parts.append(f"🔄 修改时间: {metadata['modified']}")
        
        return "\n".join(summary_parts)
    
    def _analyze_content(self, content: str, metadata: Dict, doc) -> str:
        """分析Word文档内容"""
        analysis_parts = []
        
        # 文档结构分析
        para_count = metadata.get("paragraph_count", 0)
        table_count = metadata.get("table_count", 0)
        
        if para_count > 50:
            analysis_parts.append("📚 长篇文档，内容丰富")
        elif para_count > 20:
            analysis_parts.append("📄 中等长度文档")
        else:
            analysis_parts.append("📋 简短文档")
        
        # 表格分析
        if table_count > 0:
            analysis_parts.append(f"📊 包含 {table_count} 个表格，可能包含结构化数据")
        
        # 图片分析
        if metadata.get("has_images"):
            analysis_parts.append("🖼️ 包含图片，可能是图文混排文档")
        
        # 标题分析（检测样式）
        heading_count = sum(1 for para in doc.paragraphs if para.style.name.startswith('Heading'))
        if heading_count > 0:
            analysis_parts.append(f"📑 检测到 {heading_count} 个标题，文档结构清晰")
        
        # 列表分析
        list_count = sum(1 for para in doc.paragraphs if 'List' in para.style.name)
        if list_count > 5:
            analysis_parts.append(f"📋 包含 {list_count} 个列表项")
        
        # 检测数字和统计信息
        import re
        numbers = re.findall(r'\d+', content)
        if len(numbers) > 20:
            analysis_parts.append("📈 包含大量数字数据，可能是报告或分析文档")
        
        # 检测时间信息
        years = re.findall(r'\b(19|20)\d{2}\b', content)
        if years:
            unique_years = sorted(set(years))
            if len(unique_years) > 1:
                analysis_parts.append(f"⏰ 涉及时间范围: {unique_years[0]}-{unique_years[-1]}")
        
        # 检测URL和链接
        urls = re.findall(r'https?://[^\s]+', content)
        if urls:
            analysis_parts.append(f"🔗 包含 {len(urls)} 个链接")
        
        # 检测邮箱
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', content)
        if emails:
            analysis_parts.append(f"📧 包含 {len(emails)} 个邮箱地址")
        
        return "\n".join(analysis_parts) if analysis_parts else "📄 标准Word文档"

class FileAnalysisTool:
    """文件分析工具主类"""
    
    def __init__(self, config=None):
        self.pdf_analyzer = PDFAnalyzer()
        self.table_analyzer = TableAnalyzer()
        self.docx_analyzer = DocxAnalyzer()
        self.python_analyzer = PythonCodeAnalyzer()
        self.general_code_analyzer = GeneralCodeAnalyzer()
        self.supported_types = {
            '.pdf': 'PDF文档',
            '.csv': 'CSV表格',
            '.xlsx': 'Excel表格',
            '.xls': 'Excel表格',
            '.docx': 'Word文档',
            '.doc': 'Word文档（旧版）',
            '.py': 'Python代码',
            '.java': 'Java代码',
            '.js': 'JavaScript代码',
            '.jsx': 'React代码',
            '.ts': 'TypeScript代码',
            '.tsx': 'TypeScript React代码',
            '.cpp': 'C++代码',
            '.c': 'C代码',
            '.h': 'C/C++头文件',
            '.hpp': 'C++头文件',
            '.go': 'Go代码',
            '.rs': 'Rust代码'
        }
        
        # 初始化AI Agent
        if config:
            from file_analysis_agent import FileAnalysisAgent
            self.ai_agent = FileAnalysisAgent(config)
        else:
            self.ai_agent = None
    
    def analyze_file(self, file_path: str) -> FileAnalysisResult:
        """分析文件"""
        if not os.path.exists(file_path):
            return FileAnalysisResult(
                file_type="UNKNOWN",
                file_name=os.path.basename(file_path),
                content="",
                metadata={},
                summary="",
                analysis="",
                success=False,
                error="文件不存在"
            )
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        safe_print(f"🔍 开始分析文件: {file_path} (类型: {file_ext})")
        
        if file_ext == '.pdf':
            return self.pdf_analyzer.extract_text(file_path)
        elif file_ext in ['.csv', '.xlsx', '.xls']:
            return self.table_analyzer.analyze_table(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.docx_analyzer.extract_text(file_path)
        elif file_ext == '.py':
            # Python代码分析
            code_result = self.python_analyzer.analyze(file_path)
            return self._convert_code_result(code_result)
        elif file_ext in ['.java', '.js', '.jsx', '.ts', '.tsx', '.cpp', '.c', '.h', '.hpp', '.go', '.rs']:
            # 其他编程语言代码分析
            code_result = self.general_code_analyzer.analyze(file_path)
            return self._convert_code_result(code_result)
        else:
            return FileAnalysisResult(
                file_type="UNSUPPORTED",
                file_name=os.path.basename(file_path),
                content="",
                metadata={},
                summary="",
                analysis="",
                success=False,
                error=f"不支持的文件类型: {file_ext}"
            )
    
    def _convert_code_result(self, code_result: CodeAnalysisResult) -> FileAnalysisResult:
        """将代码分析结果转换为文件分析结果"""
        return FileAnalysisResult(
            file_type=f"CODE_{code_result.language.upper()}",
            file_name=code_result.file_name,
            content=code_result.content,
            metadata={
                "language": code_result.language,
                "structure": code_result.structure,
                "metrics": code_result.metrics
            },
            summary=code_result.summary,
            analysis=code_result.analysis,
            success=code_result.success,
            error=code_result.error
        )
    
    def get_supported_types(self) -> Dict[str, str]:
        """获取支持的文件类型"""
        return self.supported_types
    
    def generate_ai_analysis(self, result: FileAnalysisResult, user_question: str = "") -> str:
        """生成AI分析报告"""
        if not result.success:
            return f"❌ 文件分析失败: {result.error}"
        
        analysis_parts = []
        
        # 基本信息
        analysis_parts.append(f"📁 **文件信息**")
        analysis_parts.append(f"- 文件名: {result.file_name}")
        analysis_parts.append(f"- 文件类型: {result.file_type}")
        analysis_parts.append("")
        
        # 摘要
        analysis_parts.append(f"📋 **文件摘要**")
        analysis_parts.append(result.summary)
        analysis_parts.append("")
        
        # 智能分析
        analysis_parts.append(f"🔍 **智能分析**")
        analysis_parts.append(result.analysis)
        analysis_parts.append("")
        
        # 使用专门的AI Agent进行深度分析
        if self.ai_agent:
            try:
                ai_analysis = self.ai_agent.analyze_file_with_ai(result, user_question)
                if ai_analysis:
                    analysis_parts.append(ai_analysis)
            except Exception as e:
                safe_print(f"⚠️ AI Agent分析失败: {e}")
                # 回退到简单分析
                ai_insights = self._generate_ai_insights(result, user_question)
                if ai_insights:
                    analysis_parts.append(f"✅ **AI深度分析**")
                    analysis_parts.append(ai_insights)
        else:
            # 回退到简单分析
            ai_insights = self._generate_ai_insights(result, user_question)
            if ai_insights:
                analysis_parts.append(f"✅ **AI深度分析**")
                analysis_parts.append(ai_insights)
        
        # 不再显示关键信息和内容预览（已包含在AI分析报告中）
        
        return "\n".join(analysis_parts)
    
    def _generate_ai_insights(self, result: FileAnalysisResult, user_question: str = "") -> str:
        """生成AI深度分析洞察"""
        try:
            # 根据文件类型生成不同的AI分析
            if result.file_type == "PDF":
                return self._analyze_pdf_with_ai(result, user_question)
            elif result.file_type == "TABLE":
                return self._analyze_table_with_ai(result, user_question)
            else:
                return self._analyze_general_with_ai(result, user_question)
        except Exception as e:
            safe_print(f"⚠️ AI分析生成失败: {e}")
            return ""
    
    def _analyze_pdf_with_ai(self, result: FileAnalysisResult, user_question: str) -> str:
        """使用AI分析PDF内容"""
        insights = []
        
        # 文档结构分析
        content = result.content
        lines = content.split('\n')
        
        # 检测章节结构
        chapters = [line.strip() for line in lines if '第' in line and '章' in line]
        if chapters:
            insights.append(f"📚 **文档结构**: 检测到 {len(chapters)} 个章节")
            insights.append(f"   主要章节: {', '.join(chapters[:3])}{'...' if len(chapters) > 3 else ''}")
        
        # 关键概念提取
        key_concepts = self._extract_key_concepts(content)
        if key_concepts:
            insights.append(f"🔑 **核心概念**: {', '.join(key_concepts[:5])}")
        
        # 数据统计
        numbers = [word for word in content.split() if word.isdigit()]
        if len(numbers) > 5:
            insights.append(f"📊 **数据丰富度**: 包含 {len(numbers)} 个数字，数据详实")
        
        # 时间线检测
        time_indicators = ['年', '年代', '世纪', '月', '日']
        time_mentions = [line for line in lines if any(indicator in line for indicator in time_indicators)]
        if time_mentions:
            insights.append(f"⏰ **时间维度**: 包含历史发展脉络")
        
        # 表格检测
        table_indicators = ['|', '\t', '  ']
        table_lines = [line for line in lines if any(indicator in line for indicator in table_indicators)]
        if len(table_lines) > 3:
            insights.append(f"📋 **结构化数据**: 包含表格或列表信息")
        
        # 智能总结生成
        smart_summary = self._generate_smart_summary(content, user_question)
        if smart_summary:
            insights.append(f"💡 **智能总结**: {smart_summary}")
        
        return "\n".join(insights) if insights else "📄 标准文档内容，结构清晰"
    
    def _generate_smart_summary(self, content: str, user_question: str = "") -> str:
        """生成智能总结"""
        try:
            # 如果内容太长，先进行文本分割
            if len(content) > 2000:
                # 使用LangChain的文本分割器
                from langchain.text_splitter import RecursiveCharacterTextSplitter
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200,
                    length_function=len,
                )
                chunks = text_splitter.split_text(content)
                
                # 取前几个chunk进行总结
                summary_content = "\n".join(chunks[:3])
            else:
                summary_content = content
            
            # 生成总结提示词
            if user_question:
                prompt = f"""
请根据用户问题"{user_question}"，对以下文档内容进行智能分析和总结：

文档内容：
{summary_content[:1500]}

请提供：
1. 核心观点总结（2-3个要点）
2. 关键数据或事实
3. 与用户问题相关的重点内容

总结要简洁明了，突出重点。
"""
            else:
                prompt = f"""
请对以下文档内容进行智能总结：

文档内容：
{summary_content[:1500]}

请提供：
1. 文档主题和核心观点
2. 主要章节或结构
3. 关键信息要点

总结要简洁明了，突出重点。
"""
            
            # 这里可以调用AI API，现在先用简单的文本分析
            return self._simple_content_analysis(summary_content)
            
        except Exception as e:
            safe_print(f"⚠️ 智能总结生成失败: {e}")
            return ""
    
    def _simple_content_analysis(self, content: str) -> str:
        """简单的内容分析（不使用AI API）"""
        try:
            lines = content.split('\n')
            
            # 提取标题和关键句
            key_sentences = []
            
            # 查找包含关键词的句子
            important_keywords = ['人工智能', 'AI', '技术', '发展', '应用', '挑战', '未来', '智能']
            
            for line in lines:
                line = line.strip()
                if len(line) > 20 and any(keyword in line for keyword in important_keywords):
                    key_sentences.append(line)
            
            # 限制句子数量
            key_sentences = key_sentences[:3]
            
            if key_sentences:
                return " | ".join([sentence[:50] + "..." if len(sentence) > 50 else sentence for sentence in key_sentences])
            else:
                # 如果没有找到关键词，返回前几句
                non_empty_lines = [line.strip() for line in lines if line.strip()]
                if non_empty_lines:
                    first_sentences = non_empty_lines[:2]
                    return " | ".join([sentence[:50] + "..." if len(sentence) > 50 else sentence for sentence in first_sentences])
                else:
                    return "文档内容已解析，包含丰富信息"
                    
        except Exception as e:
            safe_print(f"⚠️ 内容分析失败: {e}")
            return "文档内容分析完成"
    
    def _analyze_table_with_ai(self, result: FileAnalysisResult, user_question: str) -> str:
        """使用AI分析表格内容"""
        insights = []
        metadata = result.metadata
        
        # 数据规模分析
        rows = metadata.get('rows', 0)
        cols = metadata.get('columns', 0)
        insights.append(f"📊 **数据规模**: {rows} 行 × {cols} 列")
        
        # 数据类型分析
        data_types = metadata.get('data_types', {})
        type_counts = {}
        for dtype in data_types.values():
            type_name = str(dtype)
            type_counts[type_name] = type_counts.get(type_name, 0) + 1
        
        if type_counts:
            insights.append(f"🔢 **数据类型**: {', '.join([f'{dtype}({count})' for dtype, count in type_counts.items()])}")
        
        # 数据质量分析
        missing_values = sum(1 for col, count in metadata.get('missing_values', {}).items() if count > 0)
        if missing_values == 0:
            insights.append("✅ **数据质量**: 无缺失值，数据完整")
        else:
            insights.append(f"⚠️ **数据质量**: {missing_values} 列存在缺失值")
        
        # 业务价值分析
        column_names = metadata.get('column_names', [])
        if any('id' in col.lower() for col in column_names):
            insights.append("🆔 **业务特征**: 包含标识符字段")
        if any('date' in col.lower() or 'time' in col.lower() for col in column_names):
            insights.append("⏰ **时间特征**: 包含时间相关字段")
        if any('price' in col.lower() or 'cost' in col.lower() or 'amount' in col.lower() for col in column_names):
            insights.append("💰 **财务特征**: 包含金额相关字段")
        
        return "\n".join(insights) if insights else "📊 标准表格数据"
    
    def _analyze_general_with_ai(self, result: FileAnalysisResult, user_question: str) -> str:
        """使用AI分析一般文件内容"""
        return "📄 文件内容分析完成"
    
    def _extract_key_concepts(self, content: str) -> list:
        """提取关键概念"""
        # 简单的关键词提取
        import re
        
        # 提取可能的专业术语（大写字母开头的词）
        concepts = re.findall(r'\b[A-Z][a-z]+\b', content)
        
        # 提取可能的缩写词
        abbreviations = re.findall(r'\b[A-Z]{2,}\b', content)
        
        # 合并并去重
        all_concepts = list(set(concepts + abbreviations))
        
        # 过滤掉常见的非专业词汇
        common_words = {'The', 'This', 'That', 'With', 'From', 'They', 'There', 'These', 'Those'}
        filtered_concepts = [concept for concept in all_concepts if concept not in common_words]
        
        return filtered_concepts[:10]  # 返回前10个
    
    def _extract_key_points(self, result: FileAnalysisResult) -> str:
        """提取关键信息点"""
        key_points = []
        
        if result.file_type == "PDF":
            # PDF关键信息提取
            content = result.content
            
            # 提取数字信息
            import re
            numbers = re.findall(r'\d+', content)
            if numbers:
                key_points.append(f"📊 包含 {len(numbers)} 个数字数据")
            
            # 提取章节信息
            chapters = re.findall(r'第[一二三四五六七八九十\d]+章[：:]\s*([^\n]+)', content)
            if chapters:
                key_points.append(f"📚 主要章节: {', '.join(chapters[:3])}")
            
            # 提取时间信息
            years = re.findall(r'\b(19|20)\d{2}\b', content)
            if years:
                key_points.append(f"⏰ 涉及年份: {', '.join(sorted(set(years))[:5])}")
        
        elif result.file_type == "TABLE":
            # 表格关键信息提取
            metadata = result.metadata
            rows = metadata.get('rows', 0)
            cols = metadata.get('columns', 0)
            
            key_points.append(f"📊 数据规模: {rows} 行 × {cols} 列")
            
            column_names = metadata.get('column_names', [])
            if column_names:
                key_points.append(f"📝 主要字段: {', '.join(column_names[:5])}")
        
        return "\n".join(key_points) if key_points else "📄 文件内容已成功解析"

# 测试函数
def test_file_analysis():
    """测试文件分析功能"""
    tool = FileAnalysisTool()
    
    print("🧪 测试文件分析工具...")
    print(f"支持的文件类型: {tool.get_supported_types()}")
    
    # 这里可以添加测试文件路径
    # result = tool.analyze_file("test.pdf")
    # print(tool.generate_ai_analysis(result))

if __name__ == "__main__":
    test_file_analysis()
